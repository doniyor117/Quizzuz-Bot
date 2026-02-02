from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import tempfile

async def generate_set_docx(set_data: dict, cards_list: list) -> str:
    """
    Generate a DOCX file for a flashcard set.
    
    Args:
        set_data: Dictionary with set info (set_name, created_at, etc.)
        cards_list: List of card dictionaries with 'term' and 'definition'
    
    Returns:
        File path to the generated DOCX file
    """
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading(set_data.get('set_name', 'Flashcard Set'), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Created: {set_data.get('created_at', 'N/A')}\n").italic = True
    meta.add_run(f"Total Cards: {len(cards_list)}").italic = True
    
    # Add separator
    doc.add_paragraph('_' * 60)
    doc.add_paragraph()  # Empty line
    
    # Add cards
    for idx, card in enumerate(cards_list, 1):
        # Card number and term
        term_para = doc.add_paragraph()
        term_run = term_para.add_run(f"{idx}. {card.get('term', 'N/A')}")
        term_run.bold = True
        term_run.font.size = Pt(12)
        term_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
        
        # Definition
        def_para = doc.add_paragraph()
        def_para.paragraph_format.left_indent = Inches(0.5)
        def_run = def_para.add_run(card.get('definition', 'N/A'))
        def_run.font.size = Pt(11)
        
        # Add spacing between cards
        if idx < len(cards_list):
            doc.add_paragraph()
    
    # Add footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    footer = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)
    
    # Save to temporary file
    exports_dir = '/tmp/exports'
    os.makedirs(exports_dir, exist_ok=True)
    
    # Generate unique filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    set_name_safe = "".join(c for c in set_data.get('set_name', 'set') if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"{set_name_safe}_{timestamp}.docx"
    filepath = os.path.join(exports_dir, filename)
    
    # Save document
    doc.save(filepath)
    
    return filepath

async def generate_set_pdf(set_data: dict, cards_list: list) -> str:
    """
    Generate a PDF file for a flashcard set.
    
    Args:
        set_data: Dictionary with set info (set_name, created_at, etc.)
        cards_list: List of card dictionaries with 'term' and 'definition'
    
    Returns:
        File path to the generated PDF file
    """
    from fpdf import FPDF
    
    class FlashcardPDF(FPDF):
        def header(self):
            # Fonts are already added globally
            self.set_font('NotoSans', 'B', 16)
            self.cell(0, 10, set_data.get('set_name', 'Flashcard Set'), 0, 1, 'C')
            self.set_font('NotoSans', 'I', 10)
            self.cell(0, 5, f"Created: {set_data.get('created_at', 'N/A')}", 0, 1, 'C')
            self.cell(0, 5, f"Total Cards: {len(cards_list)}", 0, 1, 'C')
            self.ln(5)
        
        def footer(self):
            self.set_y(-15)
            self.set_font('NotoSans', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}/{{nb}}  |  Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', 0, 0, 'C')
    
    # Create PDF
    pdf = FlashcardPDF()
    
    # Add fonts globally first
    pdf.add_font('NotoSans', '', 'assets/fonts/NotoSans-Regular.ttf')
    pdf.add_font('NotoSans', 'B', 'assets/fonts/NotoSans-Bold.ttf')
    pdf.add_font('NotoSans', 'I', 'assets/fonts/NotoSans-Italic.ttf')
    
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Add separator
    pdf.set_draw_color(0, 0, 128)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)
    
    # Add cards
    for idx, card in enumerate(cards_list, 1):
        # Card number and term
        pdf.set_font('NotoSans', 'B', 12)
        pdf.set_text_color(0, 0, 128)  # Dark blue
        pdf.cell(0, 8, f"{idx}. {card.get('term', 'N/A')}", 0, 1)
        
        # Definition
        pdf.set_font('NotoSans', '', 11)
        pdf.set_text_color(0, 0, 0)  # Black
        pdf.set_x(20)  # Indent
        
        # Handle multi-line definitions
        definition = card.get('definition', 'N/A')
        definition = definition.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, definition)
        
        # Add spacing
        pdf.ln(4)
        
        # Check if we need a new page
        if pdf.get_y() > 250:
            pdf.add_page()
    
    # Save to temporary file
    exports_dir = '/tmp/exports'
    os.makedirs(exports_dir, exist_ok=True)
    
    # Generate unique filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    set_name_safe = "".join(c for c in set_data.get('set_name', 'set') if c.isalnum() or c in (' ', '-', '_')).strip()
    filename = f"{set_name_safe}_{timestamp}.pdf"
    filepath = os.path.join(exports_dir, filename)
    
    # Save PDF
    pdf.output(filepath)
    
    return filepath

def cleanup_export_file(filepath: str):
    """Delete the exported file after sending."""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
    except Exception as e:
        print(f"Failed to cleanup export file {filepath}: {e}")
